from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x4A, 0x7C)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x3D, 0x5A, 0x8C)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(3)

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

def bold_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True

# ===== TITLE =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('亚马逊跨境电商运营助理（管培生）\n面试准备手册')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('罗楚婷 · 2026年7月')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ===== PART 1 =====
h2('一、岗位核心拆解')

h3('1.1 这个岗位到底要什么人？')
body('表面要求：零基础可培养、英语四级、对AI感兴趣。')
body('实际要的人：一个有学习能力、有数据意识、能扛事、想赚钱的年轻人。')
body('')
body('对照你自己：')
bullet('英语六级 ✅ — 超过"四级优先"的基本线')
bullet('AI实操经验 ✅ — 独立用AI做产品，这是JD明确写的"优先录用"条件')
bullet('数据复盘能力 ✅ — 124所高校16维数据整理+评分算法，底层逻辑=运营数据分析')
bullet('学习能力 ✅ — 跨学科教学、零基础vibe coding做产品，比学亚马逊运营难多了')
bullet('目标感 ✅ — 三战考研的坚持，院长奖学金(1%)，项目全自驱推动')

h3('1.2 管培生 vs 普通运营的区别')
body('这个岗位不是招一个"干活的"，是招一个"未来能管人的"。所以面试官不会问你"会不会操作亚马逊后台"，会问你"给你一个问题你能不能自己想办法解决"。')

bullet('普通运营：学会操作→重复执行→拿提成')
bullet('管培生：轮岗学习→理解业务全貌→培养管理思维→未来带团队')

h3('1.3 五个轮岗方向解读')
body('① 全链路运营：选品→上架→Listing优化→定价→促销，了解亚马逊卖货的完整链条')
body('② 数据能力：看销量/排名/流量/转化率→找问题→提建议→优化策略→这是管理层基本功')
body('③ 市场策略：盯竞品动态→看海外趋势→制定增量方案→培养"操盘手"视角')
body('④ 风控统筹：合规+库存管理+防断货/滞销→学习全局运维思维')
body('⑤ 管理积累：跟主管参与日常管理→学协作/统筹/项目推进→为晋升打基础')

doc.add_page_break()

# ===== PART 2 =====
h2('二、亚马逊运营认知框架（零基础急救）')

h3('2.1 亚马逊是怎么卖货的？')
body('简单理解：在亚马逊上开店→上产品→优化产品页面→打广告引流→客户下单→发货→收钱')
body('')
body('核心公式：销售额 = 流量 × 转化率 × 客单价')
body('运营每天做的事，就是提高这三个变量中的每一个。')

h3('2.2 亚马逊运营的核心环节')
body('① 选品：市场调研→看什么好卖→分析竞争→决定做什么产品')
body('② 上架(Listing)：标题+图片+五点描述+A+页面→让客户看了想买')
body('③ 定价：竞品什么价→我什么价→成本多少→利润多少')
body('④ 广告/引流：站内广告(PPC)+站外引流→花钱买曝光')
body('⑤ 转化优化：看数据→哪些词转化高→哪些图点击好→持续调')
body('⑥ 库存管理：备多少货→什么时候补→不断货不滞销')
body('⑦ 账号风控：合规运营→不踩平台红线→保账号安全')

h3('2.3 运营的日常工作（一天长什么样）')
body('上午：看昨日销售数据→出单量/广告花费/转化率→有没有异常')
body('上午：检查Listing→有没有被跟卖/差评/排名下降')
body('下午：优化广告→调出价/加关键词/否词')
body('下午：看竞品→对手有没有降价/上新/换图')
body('傍晚：处理邮件/售后/补货计划')
body('下班前：复盘今天数据→记下来→明天对照')

doc.add_page_break()

h3('2.4 核心专业术语速查表')

terms = [
    ('Listing', '产品页面。包含标题、图片、五点描述、A+页面、搜索词等'),
    ('SKU', '库存单位编码。每个产品一个唯一编号，用于库存管理'),
    ('ASIN', '亚马逊标准识别号。每个上架商品自动获得的唯一ID'),
    ('FBA', '亚马逊物流。货发到亚马逊仓库，由亚马逊负责配送和客服'),
    ('FBM', '卖家自发货。自己打包寄给客户，不用亚马逊仓库'),
    ('PPC', '按点击付费广告。亚马逊站内广告，按点击收费'),
    ('CPC', '每次点击成本。广告每被点一次花多少钱'),
    ('CTR', '点击率。曝光中被点击的比例 = 点击/曝光'),
    ('CVR', '转化率。点击中实际购买的比例 = 订单/点击'),
    ('ACOS', '广告销售成本比。广告费÷广告带来的销售额，越低越好'),
    ('ROI', '投资回报率。赚的钱÷投入的钱'),
    ('ROAS', '广告支出回报。广告带来的销售额÷广告费'),
    ('BSR', '畅销榜排名。数字越小越畅销'),
    ('关键词排名', '搜索某个词时你的产品排第几位。排名越靠前流量越大'),
    ('跟卖', '其他卖家在你的Listing下以更低价格卖同样的产品'),
    ('Buy Box', '购物车。同一个产品多个卖家，亚马逊选一个给购物车'),
    ('A+页面', '品牌注册后才有的图文详情页，比普通描述更丰富'),
    ('五点描述', 'Listing中的5个核心卖点，客户第一眼看的地方'),
    ('长尾关键词', '搜索量小但精准的关键词组合，转化率通常更高'),
    ('黄金购物车', 'Buy Box的另一种说法，拿到购物车=拿到主要流量'),
    ('Review', '产品评论。客户购买后的评价，星级+文字'),
    ('Feedback', '店铺反馈。对卖家服务/物流的评价，不同于产品Review'),
    ('ACOS', '广告花费÷广告销售额。比如花10元广告卖了100元→ACOS=10%'),
]

for term, defn in terms:
    p = doc.add_paragraph()
    run = p.add_run(f'{term}：')
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(defn)
    run2.font.size = Pt(10)

doc.add_page_break()

# ===== PART 3 =====
h2('三、面试问答预演')

h3('Q1：你为什么想进入跨境电商行业？')
body('"我对跨境电商的兴趣来自两个方向。第一，我自己做过AI产品并且通过内容运营获得了3000多用户和14万曝光，这让我意识到"好产品+好运营=增长"的逻辑，而跨境电商就是这个逻辑最直接的行业——选品就是选产品、Listing优化就是内容运营、广告投放就是用户增长。第二，你们是管培生体系，轮岗学习→独立运营→管理晋升，这正是我想要的长线成长路径，而不是一个重复执行的岗位。"')

h3('Q2：你完全没做过亚马逊，凭什么觉得能胜任？')
body('"我确实没操作过亚马逊后台，但管培生岗位本身就是为新人设计的。我有三个优势可以快速上手：第一，我的学习能力是验证过的——从语文地理被调到数学化学、不会代码用AI从零做产品，我都靠自己学会了。第二，我有数据复盘的习惯——做择校工具时整理了134所学校16维数据、运营小红书时每次发完都复盘数据，这和运营亚马逊看数据找问题的底层逻辑是一样的。第三，你们的JD写了"有AI实操经验者优先"，而我可能是你们今天面试的人里AI用得最深的一个——我不只是用过，是用AI做出了上线产品。我相信这三项能力比"会点亚马逊后台按钮"更有价值。"')

h3('Q3：你怎么理解运营？')
body('"我理解的运营不是执行，是增长——销售额=流量×转化率×客单价。运营每天的事就是让这三个变量变大：用广告和排名拿流量、用好的Listing提高转化、用定价和促销拉高客单价。我过去一年做的事情本质上也是运营：我做了个产品，然后通过内容运营拿到14万曝光和3000用户——把"产品"换成"亚马逊Listing"，这完全是同一套能力模型。"')

h3('Q4：给你一个产品，你会怎么推？')
body('"如果是零基础产品，我会分四步走。第一步市场调研：看竞品的价格、销量、Review、关键词，搞清楚这个品类什么价位好卖、什么卖点客户在乎。第二步Listing优化：标题埋关键词、五点打核心卖点、图片让客户一眼看懂产品价值。第三步引流测试：先开一组自动广告跑几天，看哪些词转化好，然后手动精准投放。第四步复盘迭代：每周看数据——流量涨了没、转化率高了没、ACOS在合理范围内吗——哪里不好调哪里。"')

h3('Q5：你怎么看待压力和加班？')
body('"跨境电商有时差，旺季忙是常态。我考研坚持了三年、教培带课期间同时备考还要备课、做AI产品的时候熬夜调试是常态——我能扛压力。我追求的是付出和回报成正比，你们的薪资机制是纯业绩导向，多劳多得正是我想要的。"')

h3('Q6：你未来三年的规划？')
body('"短期：第一年按照管培生计划轮岗学习，把亚马逊全链路吃透。中期：第二年能独立负责一个品线或店铺，用业绩证明自己。长期：第三年带团队，从自己做事变成帮别人做事。我来面试管培生而不是普通运营，就是想要这条清晰的晋升路径。"')

doc.add_page_break()

# ===== PART 4 =====
h2('四、零经验新人·运营流程速成')

h3('4.1 亚马逊卖货全流程（一张图看懂）')
steps = [
    '市场调研：看什么品类好卖→竞争大不大→利润空间多少→确定做什么',
    '选品立项：找到具体产品→分析竞品Review→找出差评中的优化机会→确定差异化方向',
    '采购备货：联系供应商→打样确认→下单生产→贴标→发货到亚马逊仓库(FBA)',
    'Listing上架：写标题(埋关键词)→做图片(主图+场景图+对比图)→写五点描述→A+页面',
    '定价策略：竞品价格→我定多少→扣除成本(采购+物流+佣金+广告)→利润多少',
    '推广引流：开PPC广告→做关键词→站外引流(社交媒体/红人测评)→参加促销活动',
    '数据复盘：每日看销量/排名/广告数据→每周分析转化率和ACOS→每月做运营总结',
    '持续优化：根据数据调Listing→调广告→调价格→补货→处理售后→维护Review',
]
for s in steps:
    bullet(s)

h3('4.2 新人最该关注的三个指标')
body('每天上班第一件事看三个数：')
bullet('① 订单量：今天出了多少单？跟前天/上周同期比涨了还是跌了？')
bullet('② 广告花费和ACOS：花了多少广告费？每卖100块花了多少广告费？超过25%就要优化')
bullet('③ BSR排名：产品在大类/小类的畅销榜排名涨了还是跌了？')

h3('4.3 面试时可以说"我已经在了解的"')
body('面试官不会考你亚马逊操作细节。你只需要展现"我已经开始自学了"的态度。可以说：')
bullet('"我了解到亚马逊运营核心看三个指标：流量、转化率、ACOS"')
bullet('"我知道Listing优化最重要的是标题关键词和主图"')
bullet('"FBA是亚马逊物流，FBM是自发货，大部分卖家首选FBA"')
bullet('"广告要先自动跑数据，再手动精准投，不能一上来就精准"')

doc.add_page_break()

# ===== PART 5 =====
h2('五、你的差异化优势·面试杀手锏')

body('你不该以"零基础新人"的身份面试，而是以"有AI能力+数据思维+内容经验的新人"身份。')
body('')
bold_body('杀手锏一：AI实操经验（直接命中JD加分项）')
body('"我不只是用过AI，我是用AI独立做了产品——从调研到设计到开发到上线，全部AI辅助完成。贵司有自研运营管理系统，我对AI工具的理解和学习速度，能让我更快上手你们的系统。"')
body('')
bold_body('杀手锏二：数据复盘能力')
body('"我做过134所高校16维数据的收集和分析、搭建过评分算法、运营小红书每次发完都复盘数据。运营亚马逊的核心能力之一就是看数据找问题，而这个能力我已经在别的项目里反复训练过了。"')
body('')
bold_body('杀手锏三：内容运营=Listing优化')
body('"Listing优化的本质是内容运营——好的标题让人点进来、好的图片让人停下来、好的描述让人下单。我运营小红书1500粉丝、产出爆款笔记的经验，可以直接迁移到Listing文案和图片策略上。"')
body('')
bold_body('杀手锏四：学习能力证明')
body('"从文科生到用AI做开发、从语文老师到数学化学老师、从小白到小红书1500粉——我做的事情都有一个共同点：从零开始，快速学会，然后拿结果。"')

doc.add_page_break()

# ===== PART 6 =====
h2('六、反问面试官（展现你思考深度）')

bullet('"管培生第一年的轮岗周期大概多长？每个板块有没有阶段性的考核标准？"')
bullet('"目前团队的管培生从入职到独立负责一个品线，大概需要多长时间？"')
bullet('"公司自研的管理系统，和市面上常用的ERP工具相比最大的差异点是什么？"')
bullet('"团队目前主要做什么品类？一个人一般同时管多少个SKU？"')

bullets_fast = [
    '仪表盘看数据→今天出单/广告/排名有没有异常',
    '检查Listing→差评/跟卖/排名/购物车有没有异常',
    '优化广告→调出价/加否定词/否无效词',
    '处理站内信/售后/邮件',
    '看竞品→价格变动/上新/Review变化',
    '补货计划→库存够不够/要不要下单',
    '下班前简单记录→今天做了什么/明天重点是什么',
]

# ===== FINAL =====
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 祝你面试顺利 —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

out = 'D:/亚马逊运营面试-完整准备.docx'
doc.save(out)
print(f'Generated: {out}')
